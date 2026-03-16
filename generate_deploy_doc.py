from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('SpeechEnhanceApp — Deployment Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Azure VM + GitHub Actions CI/CD', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

def add_section(title, level=1):
    doc.add_heading(title, level=level)

def add_step(label, content):
    p = doc.add_paragraph()
    run = p.add_run(label + ': ')
    run.bold = True
    p.add_run(content)

def add_code(text):
    p = doc.add_paragraph(text, style='No Spacing')
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    p.paragraph_format.left_indent = Pt(18)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('Note: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
    p.add_run(text)

# ── OVERVIEW ──────────────────────────────────────────────────────────────────
add_section('1. Architecture Overview')
doc.add_paragraph(
    'The app is deployed on a single Azure Linux VM. GitHub Actions handles CI/CD: '
    'every push to main triggers a build, test, and deploy pipeline. '
    'Gunicorn serves the Flask app on localhost:5000. Nginx acts as a reverse proxy on port 80.'
)
doc.add_paragraph()

# ── GITHUB REPO ───────────────────────────────────────────────────────────────
add_section('2. GitHub Repository Setup')

add_section('2.1 Push code to GitHub', level=2)
add_code('git remote add origin https://github.com/youruser/SpeechEnhanceApp.git')
add_code('git push -u origin main')

add_section('2.2 .gitignore rules (important)', level=2)
doc.add_paragraph('The following are excluded from git to avoid committing large or sensitive files:')
for item in ['*.onnx  (model weights)', '*.wav   (audio files)', '*.pem   (SSH keys)', 'venv/   (virtual environment)']:
    doc.add_paragraph(item, style='List Bullet')

add_section('2.3 GitHub Actions Secrets', level=2)
doc.add_paragraph('Add these 3 secrets at: GitHub → repo → Settings → Secrets → Actions → Repository secrets')
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Secret Name'
hdr[1].text = 'Value'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True
rows_data = [
    ('VM_HOST', 'Azure VM public IP address'),
    ('VM_USER', 'azureuser'),
    ('VM_SSH_KEY', 'Full contents of .pem private key file (including BEGIN/END lines)'),
]
for i, (k, v) in enumerate(rows_data, 1):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
doc.add_paragraph()

# ── CI/CD PIPELINE ────────────────────────────────────────────────────────────
add_section('3. CI/CD Pipeline (.github/workflows/deploy.yml)')
doc.add_paragraph('Three sequential jobs run on every push to main:')
for step in [
    'build  — installs Python 3.10 and all pip dependencies',
    'test   — runs pytest tests/ -v (deploy blocked if tests fail)',
    'deploy — SSH into VM, git reset --hard origin/main, restart service',
]:
    doc.add_paragraph(step, style='List Number')

add_note('Model files (.onnx) are NOT part of the pipeline — they live on the VM permanently.')
doc.add_paragraph()

# ── AZURE VM SETUP ────────────────────────────────────────────────────────────
add_section('4. Azure VM Initial Setup (one-time)')

add_section('4.1 Create the VM', level=2)
for item in [
    'OS: Ubuntu 22.04 LTS',
    'Size: Standard B2s (2 vCPU, 4 GB RAM) minimum',
    'Open inbound ports: 22 (SSH), 80 (HTTP), 443 (HTTPS)',
]:
    doc.add_paragraph(item, style='List Bullet')

add_section('4.2 Install system packages', level=2)
add_code('sudo apt update && sudo apt install python3-pip python3-venv nginx git -y')
add_code('sudo add-apt-repository ppa:deadsnakes/ppa -y')
add_code('sudo apt update && sudo apt install python3.10 python3.10-venv -y')
add_note('Python 3.10 is required — numba==0.58.1 does not support Python 3.12.')

add_section('4.3 Clone the repository', level=2)
add_code('git clone https://github.com/youruser/SpeechEnhanceApp.git')
add_code('cd SpeechEnhanceApp')

add_section('4.4 Create virtual environment', level=2)
add_code('python3.10 -m venv venv')
add_code('source venv/bin/activate')
add_code('pip install -r requirements.txt')

add_section('4.5 Copy ONNX model files (run locally)', level=2)
doc.add_paragraph('Run this from your local machine — models are never stored in git:')
add_code('scp -i speechapp_key.pem src/models/denoise/denoise_gan.onnx \\')
add_code('    azureuser@<VM_IP>:~/SpeechEnhanceApp/src/models/denoise/')
add_code('scp -i speechapp_key.pem src/models/upsample/upsample16_gan_500.onnx \\')
add_code('    src/models/upsample/upsample48_gan_500.onnx \\')
add_code('    azureuser@<VM_IP>:~/SpeechEnhanceApp/src/models/upsample/')

add_section('4.6 Add VM deploy key to GitHub', level=2)
doc.add_paragraph('On the VM, generate an SSH key and add it as a GitHub Deploy Key:')
add_code('ssh-keygen -t ed25519 -C "speechenhance" -f ~/.ssh/id_ed25519 -N ""')
add_code('cat ~/.ssh/id_ed25519.pub')
doc.add_paragraph('Go to: GitHub → repo → Settings → Deploy keys → Add deploy key → paste the output above.')

# ── SYSTEMD SERVICE ────────────────────────────────────────────────────────────
add_section('5. Systemd Service (gunicorn auto-start)')

add_section('5.1 Create service file', level=2)
add_code('sudo nano /etc/systemd/system/speechenhance.service')
doc.add_paragraph('Contents:')
for line in [
    '[Unit]',
    'Description=SpeechEnhanceApp',
    'After=network.target',
    '',
    '[Service]',
    'User=azureuser',
    'WorkingDirectory=/home/azureuser/SpeechEnhanceApp',
    'Environment="PATH=/home/azureuser/SpeechEnhanceApp/venv/bin"',
    'Environment="SECRET_KEY=<your-generated-secret-key>"',
    'ExecStart=/home/azureuser/SpeechEnhanceApp/venv/bin/gunicorn --workers 2 --bind 127.0.0.1:5000 app:app',
    'Restart=always',
    '',
    '[Install]',
    'WantedBy=multi-user.target',
]:
    add_code(line)

add_section('5.2 Generate SECRET_KEY', level=2)
add_code('python3 -c "import secrets; print(secrets.token_hex(32))"')
doc.add_paragraph('Copy the output and paste it as the SECRET_KEY value in the service file.')

add_section('5.3 Enable and start', level=2)
add_code('sudo systemctl daemon-reload')
add_code('sudo systemctl enable speechenhance')
add_code('sudo systemctl start speechenhance')
add_code('sudo systemctl status speechenhance')

add_section('5.4 Allow CI/CD to restart without password', level=2)
add_code('sudo visudo')
doc.add_paragraph('Add at the bottom:')
add_code('azureuser ALL=(ALL) NOPASSWD: /bin/systemctl restart speechenhance')
doc.add_paragraph()

# ── NGINX ─────────────────────────────────────────────────────────────────────
add_section('6. Nginx Reverse Proxy')

add_section('6.1 Create config', level=2)
add_code('sudo nano /etc/nginx/sites-available/speechenhance')
doc.add_paragraph('Contents:')
for line in [
    'server {',
    '    listen 80;',
    '    server_name <VM_IP_OR_DOMAIN>;',
    '',
    '    client_max_body_size 50M;',
    '',
    '    location / {',
    '        proxy_pass http://127.0.0.1:5000;',
    '        proxy_set_header Host $host;',
    '        proxy_set_header X-Real-IP $remote_addr;',
    '    }',
    '}',
]:
    add_code(line)

add_section('6.2 Enable and restart', level=2)
add_code('sudo ln -s /etc/nginx/sites-available/speechenhance /etc/nginx/sites-enabled/')
add_code('sudo rm /etc/nginx/sites-enabled/default')
add_code('sudo nginx -t')
add_code('sudo systemctl restart nginx')
doc.add_paragraph()

# ── DOMAIN + CLOUDFLARE ───────────────────────────────────────────────────────
add_section('7. Custom Domain with Cloudflare + HTTPS')

add_section('7.1 Buy a domain on Cloudflare', level=2)
doc.add_paragraph('Go to cloudflare.com → Domain Registration → Register Domains. Search and purchase your domain (~$10–20/yr).')

add_section('7.2 Point DNS to the VM', level=2)
doc.add_paragraph('In Cloudflare dashboard → your domain → DNS → Records → Add record:')
table = doc.add_table(rows=3, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(['Type', 'Name', 'Content', 'Proxy']):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
table.rows[1].cells[0].text = 'A'
table.rows[1].cells[1].text = '@'
table.rows[1].cells[2].text = '74.234.27.182'
table.rows[1].cells[3].text = 'DNS only (grey cloud)'
table.rows[2].cells[0].text = 'A'
table.rows[2].cells[1].text = 'www'
table.rows[2].cells[2].text = '74.234.27.182'
table.rows[2].cells[3].text = 'DNS only (grey cloud)'
doc.add_paragraph()

add_section('7.3 Update nginx server_name', level=2)
doc.add_paragraph('Edit /etc/nginx/sites-available/speechenhance and change server_name to your domain:')
add_code('server_name yourdomain.com www.yourdomain.com;')
add_code('sudo nginx -t && sudo systemctl restart nginx')

add_section('7.4 Install free SSL certificate with Let\'s Encrypt', level=2)
add_code('sudo apt install certbot python3-certbot-nginx -y')
add_code('sudo certbot --nginx -d yourdomain.com -d www.yourdomain.com')
doc.add_paragraph('Certbot automatically updates nginx config for HTTPS and sets up auto-renewal. Certificate expires every 90 days and renews automatically.')

add_section('7.5 Enable Cloudflare proxy', level=2)
doc.add_paragraph('After certbot succeeds, turn on the Cloudflare proxy for DDoS protection:')
doc.add_paragraph('Cloudflare → DNS → Records → click grey cloud next to both A records to turn orange.', style='List Bullet')

add_section('7.6 Set Cloudflare SSL mode to Full (strict)', level=2)
doc.add_paragraph('Cloudflare → SSL/TLS → Overview → select Full (strict).')
add_note('Do NOT use "Flexible" or plain "Full" — this causes connection failures. Must be Full (strict) when the origin has a valid certificate.')
doc.add_paragraph()

add_section('7.7 Azure NSG — ensure ports are open', level=2)
doc.add_paragraph('Azure Portal → Virtual machines → speechenhance → Networking → Inbound port rules:')
table2 = doc.add_table(rows=3, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, h in enumerate(['Port', 'Protocol', 'Name']):
    hdr2[i].text = h
    hdr2[i].paragraphs[0].runs[0].bold = True
table2.rows[1].cells[0].text = '80'
table2.rows[1].cells[1].text = 'TCP'
table2.rows[1].cells[2].text = 'HTTP'
table2.rows[2].cells[0].text = '443'
table2.rows[2].cells[1].text = 'TCP'
table2.rows[2].cells[2].text = 'HTTPS'
doc.add_paragraph()

# ── SECURITY NOTE ─────────────────────────────────────────────────────────────
add_section('8. Security Notes')
for item in [
    'Never commit .pem files — they are in .gitignore as *.pem',
    'If a .pem is accidentally pushed: revoke the key on Azure immediately, remove from git history with git filter-branch, force push, generate a new key',
    'SECRET_KEY must be set as an environment variable — never hardcoded in source',
    'Model files are large (0.5 GB) and must never be committed to git',
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

# ── APP URL ───────────────────────────────────────────────────────────────────
add_section('9. App URL')
doc.add_paragraph('After setup the app is accessible at:')
add_code('http://<VM_PUBLIC_IP>')
doc.add_paragraph(
    'To add HTTPS with a custom domain: buy a domain, point its DNS A record to the VM IP, '
    'then run certbot on the VM for a free SSL certificate.'
)

# Save
output_path = 'SpeechEnhanceApp_Deployment_Guide.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
